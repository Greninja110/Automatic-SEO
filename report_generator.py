"""
Report Generator — Compact .docx SEO audit report.
FAIL and WARN checks shown individually with full detail.
PASS checks collapsed into a single summary row per module.
"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from checks import MODULE_TITLES

# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
C_PASS_BG   = "C6EFCE"
C_FAIL_BG   = "FFC7CE"
C_WARN_BG   = "FFEB9C"
C_PASS_TXT  = "276221"
C_FAIL_TXT  = "9C0006"
C_WARN_TXT  = "9C5700"
C_HEADER_BG = "1F3864"
C_HEADER_FG = "FFFFFF"
C_TITLE_FG  = "1F3864"
C_PAGE_FG   = "595959"
C_PASS_COLL = "EAF4EA"   # very light green for collapsed pass row

# ---------------------------------------------------------------------------
# oxml helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="D9D9D9", size="4") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_run_color(run, hex_color: str) -> None:
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )


def _page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _style_heading(paragraph, level: int = 1) -> None:
    for run in paragraph.runs:
        run.font.name  = "Calibri"
        run.font.size  = Pt(13 if level == 1 else 11)
        run.bold       = True
        _set_run_color(run, C_TITLE_FG)


def _trunc(text: str, n: int = 150) -> str:
    """Truncate long strings so cells stay compact."""
    text = str(text)
    return text if len(text) <= n else text[:n] + "…"


def _short_url(url: str, n: int = 55) -> str:
    """Show only the path+domain if URL is long."""
    return url if len(url) <= n else "…" + url[-(n-1):]


# ---------------------------------------------------------------------------
# Cover page  (compact — 1 page)
# ---------------------------------------------------------------------------

def _add_cover_page(doc: Document, metadata: dict, score: dict) -> None:
    pct = score["score_pct"]
    if pct >= 80:
        badge_color, grade = "00B050", "Good"
    elif pct >= 60:
        badge_color, grade = "FFC000", "Needs Improvement"
    else:
        badge_color, grade = "FF0000", "Poor"

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SEO AUDIT REPORT")
    r.bold = True; r.font.size = Pt(28)
    _set_run_color(r, C_TITLE_FG)

    # URL + date on one line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{metadata['url']}   |   {metadata['date']}")
    r.font.size = Pt(11)
    _set_run_color(r, "444444")

    # Score badge
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{pct:.0f}/100  {grade}")
    r.bold = True; r.font.size = Pt(36)
    _set_run_color(r, badge_color)

    # Stats in a compact 4-col table
    doc.add_paragraph()
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    hdrs   = ["Total Checks", "Passed",    "Failed",   "Warnings"]
    vals   = [score["total"], score["passed"], score["failed"], score["warnings"]]
    colors = [C_HEADER_BG,    "276221",    "9C0006",   "9C5700"]
    for i, (h, v, bg) in enumerate(zip(hdrs, vals, colors)):
        hc = tbl.rows[0].cells[i]; vc = tbl.rows[1].cells[i]
        _set_cell_bg(hc, bg); _set_cell_bg(vc, bg)
        hr = hc.paragraphs[0].add_run(h)
        hr.bold = True; hr.font.size = Pt(9)
        _set_run_color(hr, C_HEADER_FG)
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vc.paragraphs[0].add_run(str(v))
        vr.bold = True; vr.font.size = Pt(16)
        _set_run_color(vr, C_HEADER_FG)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Crawl info
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Pages crawled: {metadata['pages_crawled']}   |   "
        f"Depth: {metadata['depth']}   |   "
        f"Keyword: {metadata.get('keyword') or 'N/A'}"
    )
    r.font.size = Pt(10)
    _set_run_color(r, "666666")


# ---------------------------------------------------------------------------
# Summary + TOC  (single page — one compact table)
# ---------------------------------------------------------------------------

def _make_bar(pass_rate: float, width: int = 15) -> str:
    filled = round(pass_rate * width)
    return "█" * filled + "░" * (width - filled) + f" {pass_rate*100:.0f}%"


def _add_summary_toc(doc: Document, score: dict, all_results: dict) -> None:
    h = doc.add_heading("Summary & Module Overview", level=1)
    _style_heading(h)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    for i, (txt, bg) in enumerate([
        ("Module", C_HEADER_BG), ("FAIL", C_HEADER_BG),
        ("WARN", C_HEADER_BG),   ("PASS", C_HEADER_BG),
        ("Progress", C_HEADER_BG)
    ]):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, bg)
        r = cell.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(9)
        _set_run_color(r, C_HEADER_FG)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for module_name, result in all_results.items():
        checks   = result.get("checks", [])
        n_fail   = sum(1 for c in checks if c["status"] == "fail")
        n_warn   = sum(1 for c in checks if c["status"] == "warning")
        n_pass   = sum(1 for c in checks if c["status"] == "pass")
        total    = len(checks)
        rate     = n_pass / total if total else 0
        title    = MODULE_TITLES.get(module_name, module_name)

        row = tbl.add_row().cells

        # Module name
        r = row[0].paragraphs[0].add_run(title)
        r.font.size = Pt(9)
        if n_fail > 0:
            _set_cell_bg(row[0], "FFF5F5")

        # FAIL count
        _set_cell_bg(row[1], C_FAIL_BG if n_fail else "FFFFFF")
        r = row[1].paragraphs[0].add_run(str(n_fail))
        r.bold = (n_fail > 0); r.font.size = Pt(9)
        _set_run_color(r, C_FAIL_TXT if n_fail else "999999")
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # WARN count
        _set_cell_bg(row[2], C_WARN_BG if n_warn else "FFFFFF")
        r = row[2].paragraphs[0].add_run(str(n_warn))
        r.bold = (n_warn > 0); r.font.size = Pt(9)
        _set_run_color(r, C_WARN_TXT if n_warn else "999999")
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # PASS count
        _set_cell_bg(row[3], C_PASS_BG if n_pass == total else "FFFFFF")
        r = row[3].paragraphs[0].add_run(str(n_pass))
        r.font.size = Pt(9)
        _set_run_color(r, C_PASS_TXT)
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Progress bar
        bar = _make_bar(rate)
        r = row[4].paragraphs[0].add_run(bar)
        r.font.size = Pt(8); r.font.name = "Courier New"
        _set_run_color(r, C_PASS_TXT if rate >= 0.8 else (C_WARN_TXT if rate >= 0.5 else C_FAIL_TXT))


# ---------------------------------------------------------------------------
# Module section  (compact — grouped by check name, PASS collapsed)
# ---------------------------------------------------------------------------

def _group_checks(checks: list) -> list:
    """
    Group repeated per-page checks by (name, status).
    Returns one entry per unique check name with worst status,
    a representative detail, a count, and up to 3 sample page URLs.
    Result is sorted: FAIL first, then WARN.
    """
    from collections import OrderedDict
    groups = OrderedDict()
    for c in checks:
        key = (c.get("name", ""), c.get("status", "pass"))
        if key not in groups:
            groups[key] = {"name": c.get("name", ""), "status": c.get("status"),
                           "detail": c.get("detail", ""), "pages": [], "count": 0}
        groups[key]["count"] += 1
        page = c.get("page", "")
        if page and len(groups[key]["pages"]) < 3:
            if page not in groups[key]["pages"]:
                groups[key]["pages"].append(page)

    # For duplicate check names with different statuses, keep worst (fail > warn > pass)
    merged = {}
    STATUS_RANK = {"fail": 0, "warning": 1, "pass": 2}
    for (name, status), g in groups.items():
        if name not in merged or STATUS_RANK[status] < STATUS_RANK[merged[name]["status"]]:
            merged[name] = g
        else:
            merged[name]["count"] += g["count"]
            for p in g["pages"]:
                if p not in merged[name]["pages"] and len(merged[name]["pages"]) < 3:
                    merged[name]["pages"].append(p)

    result = sorted(merged.values(), key=lambda g: STATUS_RANK[g["status"]])
    return result


def _add_module_section(doc: Document, module_name: str, result: dict) -> None:
    title   = MODULE_TITLES.get(module_name, module_name)
    checks  = result.get("checks", [])
    summary = result.get("summary", "")

    # Heading
    h = doc.add_heading(title, level=1)
    _style_heading(h)

    # One-line summary
    p = doc.add_paragraph()
    r = p.add_run(summary)
    r.italic = True; r.font.size = Pt(9)
    _set_run_color(r, "555555")

    if not checks:
        return

    n_pass = sum(1 for c in checks if c["status"] == "pass")
    grouped = _group_checks(checks)
    actionable = [g for g in grouped if g["status"] in ("fail", "warning")]

    if not actionable:
        p = doc.add_paragraph()
        r = p.add_run(f"  ✔  All checks passed.")
        r.bold = True; r.font.size = Pt(9)
        _set_run_color(r, C_PASS_TXT)
        return

    # Table: Status | Check | Detail (with page count) | Sample Pages
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Cm(1.6), Cm(4.8), Cm(10.5), Cm(4.0)]
    for i, w in enumerate(col_widths):
        for cell in tbl.columns[i].cells:
            cell.width = w

    # Header
    for i, (txt, align) in enumerate([
        ("Status", WD_ALIGN_PARAGRAPH.CENTER),
        ("Check",  WD_ALIGN_PARAGRAPH.LEFT),
        ("Detail", WD_ALIGN_PARAGRAPH.LEFT),
        ("Pages",  WD_ALIGN_PARAGRAPH.LEFT),
    ]):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, C_HEADER_BG)
        _set_cell_border(cell)
        r = cell.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(9); r.font.name = "Calibri"
        _set_run_color(r, C_HEADER_FG)
        cell.paragraphs[0].alignment = align

    for g in actionable:
        status = g["status"]
        bg     = C_FAIL_BG  if status == "fail" else C_WARN_BG
        tc     = C_FAIL_TXT if status == "fail" else C_WARN_TXT
        label  = "FAIL"     if status == "fail" else "WARN"
        count  = g["count"]

        row = tbl.add_row().cells

        # Status
        _set_cell_bg(row[0], bg)
        _set_cell_border(row[0])
        r = row[0].paragraphs[0].add_run(label)
        r.bold = True; r.font.size = Pt(9); r.font.name = "Calibri"
        _set_run_color(r, tc)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Check name + count if > 1
        _set_cell_bg(row[1], bg + "44" if status == "fail" else "FFFEF5")
        _set_cell_border(row[1])
        name_text = g["name"] + (f" ({count}×)" if count > 1 else "")
        r = row[1].paragraphs[0].add_run(name_text)
        r.font.size = Pt(9); r.font.name = "Calibri"
        if status == "fail":
            r.bold = True

        # Detail
        _set_cell_bg(row[2], "FFFFFF")
        _set_cell_border(row[2])
        r = row[2].paragraphs[0].add_run(_trunc(g["detail"], 120))
        r.font.size = Pt(8.5); r.font.name = "Calibri"

        # Sample pages (up to 3, one per line)
        _set_cell_bg(row[3], "FFFFFF")
        _set_cell_border(row[3])
        pages = g["pages"]
        if pages:
            page_text = "\n".join(_short_url(p, 45) for p in pages)
            if count > len(pages):
                page_text += f"\n+{count - len(pages)} more"
        else:
            page_text = ""
        r = row[3].paragraphs[0].add_run(page_text)
        r.font.size = Pt(7.5); r.font.name = "Calibri"
        _set_run_color(r, C_PAGE_FG)

    # Collapsed pass row
    n_pass_grouped = sum(1 for g in grouped if g["status"] == "pass")
    if n_pass_grouped > 0:
        row = tbl.add_row().cells
        for cell in row:
            _set_cell_bg(cell, C_PASS_COLL)
            _set_cell_border(cell)
        row[0].merge(row[3])
        r = row[0].paragraphs[0].add_run(f"✔  {n_pass} check(s) passed")
        r.bold = True; r.font.size = Pt(9)
        _set_run_color(r, C_PASS_TXT)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def _add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SEO Audit Report  |  Page ")
    r.font.size = Pt(8)
    _set_run_color(r, "888888")
    for ftype, text in [("begin", None), (None, " PAGE "), ("separate", None), ("end", None)]:
        if ftype:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), ftype)
            p.runs[-1]._r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.text = text
            p.runs[-1]._r.append(it)
    r2 = p.add_run()
    r2.font.size = Pt(8)
    _set_run_color(r2, "888888")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def generate_report(all_results: dict, score: dict, metadata: dict, output_path: str) -> str:
    doc = Document()

    # A4 landscape
    sec = doc.sections[0]
    sec.page_width    = Cm(29.7)
    sec.page_height   = Cm(21.0)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    # Page 1 — Cover
    print("  Building cover page...")
    _add_cover_page(doc, metadata, score)
    _page_break(doc)

    # Page 2 — Summary + TOC
    print("  Building summary table...")
    _add_summary_toc(doc, score, all_results)
    _page_break(doc)

    # Pages 3+ — All 23 module sections (no page break between them)
    total = len(all_results)
    for i, (module_name, result) in enumerate(all_results.items(), 1):
        print(f"  Building section {i}/{total}: {module_name} ...")
        _add_module_section(doc, module_name, result)

    _add_footer(doc)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path
